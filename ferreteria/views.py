from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth import logout, authenticate, login
from .models import Boleta, Producto, Categoria, detalle_boleta, Pedido, DetallePedido
from ferreteria.carrito import Carrito
from .forms import ContactoForm, CustomUserProfileForm, ProductoForm, CustomUserCreationForm
from django.contrib import messages
from django.contrib.auth import authenticate, login
from rest_framework import viewsets
from .serializers import ProductoSerializer
from django.contrib.auth.decorators import login_required, permission_required, user_passes_test
from django.views.generic import CreateView
from django.contrib.auth.models import User
from django.core.mail import send_mail
from django.http import HttpResponseRedirect, HttpResponse
from django.urls import reverse
from django.shortcuts import render
import bcchapi
from datetime import datetime,timedelta
from docx import Document
import pandas as pd
import matplotlib.pyplot as plt
import io
from docx.shared import Inches
from django.core.paginator import Paginator
from django.http import JsonResponse
import json


# Create your views here.

class CategoriaViewset(viewsets.ModelViewSet):
    queryset = Categoria.objects.all()
    serializer_class = ProductoSerializer

class ProductoViewset(viewsets.ModelViewSet): #Componente que se encarga de guardar los datos
    queryset = Producto.objects.all()
    serializer_class = ProductoSerializer

    def get_queryset(self):
        productos = Producto.objects.all()
                         #DICCIONARIO.metodo
        nombre = self.request.GET.get('nombre')

        if nombre:
            productos = productos.filter(nombre__constains=nombre)

        return productos

class UserCreateView(CreateView):
    model = User
    form_class = CustomUserProfileForm
    template_name = 'registration/perfil.html'

def perfil(request):
    return render(request, 'registration/perfil.html')

def index(request):
    return render(request, 'index.html')

def nosotros(request):
    data = {
        'form': ContactoForm()
    }

    if request.method == 'POST':
        formulario = ContactoForm(data=request.POST)
        if formulario.is_valid():
            formulario.save()
            data["mensaje"] = "Contacto guardado"
        else:
            data["form"] = formulario
    return render (request, 'nosotros.html', data )

@login_required
def cerrar(request):
    logout(request)
    return redirect('index')

def registrar(request):
    data = {
        'form': CustomUserCreationForm()
    }
    if request.method == 'POST':
        formulario = CustomUserCreationForm(data=request.POST)
        if formulario.is_valid():
            formulario.save()
            user = authenticate(username=formulario.cleaned_data["username"], password=formulario.cleaned_data["password1"])
            login(request, user)
            messages.success(request, "Te has registrado correctamente")
            return redirect(to="index")
        data['form'] = formulario  
        
    return render(request, 'registration/registrar.html',data)

@permission_required ('ferreteria.add_producto')      
def agregar(request):

    data = {
        'form': ProductoForm()
    }

    if request.method== 'POST':
        formulario = ProductoForm(data=request.POST, files=request.FILES)
        if formulario.is_valid():
            formulario.save()
            messages.success(request, "Producto agregado")
        else:
            data["form"] = formulario
    return render(request, 'producto/agregar.html', data)

@permission_required ('ferreteria.view_producto')
def lista(request):
    productos = Producto.objects.all()
    data={
       'producto': productos
    }
    
    return render(request, 'producto/lista.html', data)

@permission_required ('ferreteria.change_producto')
def modificar(request, id):

    producto = get_object_or_404(Producto, idProducto=id)

    data = {
        'form': ProductoForm(instance=producto)
    }

    if request.method== 'POST':
        formulario = ProductoForm(data=request.POST, instance=producto, files=request.FILES)
        if formulario.is_valid():
            formulario.save()
            messages.success(request, "Modificado correctamente")
        return redirect(to="lista")
    
    return render (request, 'producto/modificar.html', data)

@login_required
def detalle_producto(request, id):
    producto = get_object_or_404(Producto, idProducto=id)   
    return render (request, 'detalleProducto.html', {'producto':producto})

@permission_required ('ferreteria.delete_producto')
def eliminar(request, id):
    producto = get_object_or_404(Producto, idProducto=id)
    producto.delete()
    messages.success(request, "Eliminado correctamente")
    return redirect(to="lista")


@permission_required ('ferreteria.view_producto')
def orden_pedidos(request):
    pedidos = Pedido.objects.all().order_by('-fecha_compra') 
    data = {}
    
    data = {'pedido':pedidos}
    return render(request, 'bodega.html', data)

@permission_required ('ferreteria.view_producto')
def detalle_pedidos(request,id):
    detalles = DetallePedido.objects.filter(id_pedido=id)
    return render (request, 'bodegaPedido.html', {'detalles':detalles})

def despachar(request, id):
    if request.method == "POST":
        pedido = get_object_or_404(Pedido,id_pedido = id)
        pedido.estado = "despachado"
        pedido.save()

    return redirect('bodega')

def admin_view(request):
    pedidos_list = Pedido.objects.all().order_by('-fecha_compra')  # Puedes ordenar como prefieras
    paginator = Paginator(pedidos_list, 10 )  # Muestra 10 pedidos por página

    page_number = request.GET.get('page')
    page_obj = paginator.get_page(page_number)

    return render(request, 'administrador.html', {'page_obj': page_obj})

def detalle_pedidosAdmin(request,id):
    detalles = DetallePedido.objects.filter(id_pedido=id)
    return render (request, 'detallepedido.html', {'detalles':detalles})

def vendedor_view(request):
    # estado= "pagado"
    estado= "pagado"
    pedidos = Pedido.objects.filter(estado=estado).order_by('-fecha_compra') 
    data = {}
    
    data = {'pedido':pedidos}
    return render(request, 'vendedor.html', data)

def detalle_pedidosVendedor(request,id):
    detalles = DetallePedido.objects.filter(id_pedido=id)
    return render (request, 'vendedorDetalles.html', {'detalles':detalles})

def entregas_view(request):
    #estado= "pagado"
    estado1= "despachado"
    estado2="entregado"
    tipo_entrega="retiro"
    pedidos = Pedido.objects.filter(estado__in=[estado1,estado2],tipo_envio=tipo_entrega).order_by('-fecha_compra') 
    data = {}
    
    data = {'pedido':pedidos}
    return render(request, 'entregas.html', data)

def entregar(request, id):
    if request.method == "POST":
        pedido = get_object_or_404(Pedido,id_pedido = id)
        pedido.estado = "entregado"
        pedido.save()

    return redirect('entregas')

def preparar(request, id):
    if request.method == "POST":
        pedido = get_object_or_404(Pedido,id_pedido = id)
        pedido.estado = "preparando"
        pedido.save()

    return redirect('vendedor')

def rechazar(request, id):
    if request.method == "POST":
        pedido = get_object_or_404(Pedido,id_pedido = id)
        pedido.estado = "rechazado"
        pedido.save()

    return redirect('vendedor')


##Los views q hizo
#VISTAS ADMIN
@user_passes_test(lambda u: u.is_superuser)
@login_required
def panel_admin_pedidos(request):
    pedidos_list = Pedido.objects.select_related('user').order_by('-fecha_compra')
    paginator = Paginator(pedidos_list, 5)  # Muestra 10 pedidos por página

    page_number = request.GET.get('page')
    page_obj = paginator.get_page(page_number)

    return render(request, 'admin/admin_panel_pedidos.html', {'page_obj': page_obj})


@user_passes_test(lambda u: u.is_superuser)
@login_required
def panel_admin_usuarios(request):
    usuarios = User.objects.all()
    return render(request, 'admin/admin_panel_usuarios.html', {'usuarios': usuarios})


@user_passes_test(lambda u: u.is_superuser)
@login_required
def actualizar_estado_pedido(request, id_pedido):
    if request.method == 'POST':
        try:
            data = json.loads(request.body)
            nuevo_estado = data.get('estado')

            pedido = Pedido.objects.get(id_pedido=id_pedido)
            pedido.estado = nuevo_estado
            pedido.save()

            return JsonResponse({'status': 'ok', 'nuevo_estado': pedido.estado})
        except Exception as e:
            return JsonResponse({'status': 'error', 'mensaje': str(e)}, status=400)

@permission_required('ferreteria.view_pedido')
def detalle_pedido_admin(request, pedido_id):
    pedido = Pedido.objects.get(pk=pedido_id)
    
    # Obtener los detalles del pedido (productos en el pedido)
    detalles_pedido = DetallePedido.objects.filter(id_pedido=pedido)

    # Obtener la boleta asociada al pedido
    boleta = pedido.boleta  # puede ser None
    
    detalles_boleta = []
    if boleta:
        # Obtener los detalles de la boleta (productos en la boleta)
        detalles_boleta = detalle_boleta.objects.filter(id_boleta=boleta)
    
    context = {
        'pedido': pedido,
        'detalles_pedido': detalles_pedido,
        'detalles_boleta': detalles_boleta,
    }
    
    return render(request, 'admin/detalle_pedido_admin.html', context)

def exportar_reporte_mensual(request):
    mes_actual =datetime.now().month
    año_actual = datetime.now().year
    #Cambiar a pagadooo!!!!!!
    estado = "pagado"
    estado2 = "entregado"
    estado3 = "despachado"
    estado4= "preparando"

    pedidos = Pedido.objects.filter(
        fecha_compra__year=año_actual,
        fecha_compra__month=mes_actual,
        estado__in = [estado,estado2,estado3,estado4]
    )

    # Convertir queryset a DataFrame
    df = pd.DataFrame.from_records(
        pedidos.values('id_pedido', 'fecha_compra', 'user__username', 'estado', 'tipo_envio', 'total'),
    )

    total_pedidos = df.shape[0]
    total_vendido = df['total'].sum()
    pedidos_entregados = df[df['estado'] == estado2].shape[0]
    pedidos_pagados = df[df['estado'] == estado].shape[0]
    iva = total_vendido * 19 / 119
    neto =  total_vendido /1.19

    # Crear el documento Word
    doc = Document()
    doc.add_heading('Informe de Ventas - Mes Actual', 0)

    doc.add_paragraph(f'Mes: {mes_actual} / Año: {año_actual}')
    doc.add_paragraph(f'Total de pedidos vendidos: {total_pedidos}')
    doc.add_paragraph(f'Total de pedidos en estado pagados: {pedidos_pagados}')
    doc.add_paragraph(f'Total de pedidos entregados: {pedidos_entregados}')
    doc.add_paragraph(f'Total vendido: ${total_vendido:,}')
    doc.add_paragraph(f'Total neto: ${neto:,.0f}')
    doc.add_paragraph(f'Iva incluido: ${iva:,.0f}')

    # Agregar tabla de pedidos
    doc.add_heading('Detalle de pedidos', level=1)
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = 'Light Grid'

    # Encabezados
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(df.columns):
        hdr_cells[i].text = col.replace('_', ' ').capitalize()

    # Filas
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)

    # Responder con archivo Word
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename=Informe_VentasFerremas_{mes_actual}_{año_actual}.docx'
    doc.save(response)
    return response


##Contador
def obtener_valor_dolar():
    user = "es.ibarra@duocuc.cl"
    password = "Uwu7832!"
    siete = bcchapi.Siete(user, password)
    today = datetime.now()

    while True:
        fecha_str = today.strftime("%Y-%m-%d")

        cuadro = siete.cuadro(
            series=["F073.TCO.PRE.Z.D"],
            nombres=["dolar"],
            desde=fecha_str,
            hasta=fecha_str,
            observado={"dolar": "last"}
        )

        if not cuadro.empty:
            valor_dolar = cuadro.iloc[0]["dolar"]
            break  
        else:
            today -= timedelta(days=1) 

    return valor_dolar

def contador_view(request):
    estado = request.GET.get('estado')

    if estado in ['pagado', 'entregado']: 
        pedidos = Pedido.objects.filter(estado=estado).order_by('-fecha_compra') 
    else:
        pedidos = Pedido.objects.all().order_by('-fecha_compra') 

    total_pedidos = pedidos.count()
    valor_dolar = obtener_valor_dolar()
    return render(request, 'contador.html', {
        'pedido': pedidos,
        'totalPedidos': total_pedidos,
        'estado_actual': estado,
        'valor_dolar': valor_dolar
    })

def contadorTrans_view(request):
    # estado= "pagado"
    estado= "pendiente"
    pedidos = Pedido.objects.filter(estado=estado).order_by('-fecha_compra')

    data = {'pedido':pedidos}
    return render(request, 'contadorTrans.html', data)

def cont_aprobar(request, id):
    if request.method == "POST":
        pedido = get_object_or_404(Pedido,id_pedido = id)
        pedido.estado = "pagado"
        pedido.save()

    return redirect('contadorTrans')

def cont_rechazar(request, id):
    if request.method == "POST":
        pedido = get_object_or_404(Pedido,id_pedido = id)
        pedido.estado = "rechazado"
        pedido.save()

    return redirect('contadorTrans')



